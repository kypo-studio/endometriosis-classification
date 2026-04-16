"""Génère le rapport au format .docx à partir du même contenu que rapport.md."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
FIG = HERE / "figures"
OUT = HERE / "rapport.docx"

doc = Document()

# --- Mise en page : marges 2 cm ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# --- Styles : police Calibri, tailles ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for lvl in range(1, 4):
    h = doc.styles[f"Heading {lvl}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

# --- Helpers ---
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p

def add_figure(filename, width_cm=15, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(FIG / filename), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)

def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
    # Rows
    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    if widths:
        for ci, w in enumerate(widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    return table

def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p

def add_numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    return p

# ================================================================
# TITRE + MÉTADONNÉES
# ================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Rapport — Classification de l'endométriose")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    "Projet : Prédiction du diagnostic d'endométriose à partir de variables cliniques\n"
    "Dataset : Endometriosis Dataset (Kaggle) — 10 000 observations synthétiques\n"
    "Date : 2026-04-16"
)
r.italic = True
r.font.size = Pt(10)

doc.add_paragraph()

# ================================================================
# RÉSUMÉ EXÉCUTIF
# ================================================================
add_heading("Résumé exécutif", level=1)
add_para(
    "Ce rapport présente une analyse exploratoire et une comparaison de 5 modèles de "
    "classification sur un jeu de données synthétique de 10 000 patientes, destiné à "
    "prédire la présence d'endométriose à partir de 6 variables cliniques (âge, BMI, "
    "douleur chronique, irrégularité menstruelle, anomalie hormonale, infertilité)."
)
add_para("Principaux résultats :", bold=True)
add_bullet("Données propres, sans valeurs manquantes ni doublons.")
add_bullet("Classes modérément déséquilibrées : 59 % de cas négatifs, 41 % de cas positifs.")
add_bullet(
    "Les variables les plus discriminantes sont l'anomalie hormonale (r = 0,19) et le "
    "niveau de douleur chronique (r = 0,12)."
)
add_bullet(
    "Le SVM (noyau RBF) obtient le meilleur F1-score sur le test set (0,577), suivi de "
    "près par la Logistic Regression (0,574)."
)
add_bullet(
    "Les performances globales restent modestes (accuracy ≈ 0,61, ROC-AUC ≈ 0,66), "
    "cohérent avec un dataset synthétique à forte composante aléatoire."
)

# ================================================================
# 1. CONTEXTE
# ================================================================
add_heading("1. Contexte et objectifs", level=1)
add_para(
    "L'endométriose est une maladie gynécologique chronique qui touche environ une femme "
    "sur dix en âge de procréer. Son diagnostic est souvent tardif (plusieurs années en "
    "moyenne), ce qui motive le développement d'outils d'aide au dépistage basés sur les "
    "symptômes rapportés."
)
add_para("Objectifs du projet :", bold=True)
add_numbered("Explorer les distributions et corrélations du dataset.")
add_numbered("Entraîner et comparer plusieurs modèles de classification supervisée.")
add_numbered("Identifier les variables les plus informatives.")

# ================================================================
# 2. DATASET
# ================================================================
add_heading("2. Présentation du dataset", level=1)

add_heading("2.1 Variables", level=2)
add_table(
    ["Variable", "Type", "Plage / Modalités", "Description"],
    [
        ["Age", "entier", "18–49", "Âge de la patiente"],
        ["Menstrual_Irregularity", "binaire", "0 / 1", "Irrégularité menstruelle"],
        ["Chronic_Pain_Level", "continu", "0–10", "Niveau de douleur chronique"],
        ["Hormone_Level_Abnormality", "binaire", "0 / 1", "Anomalie des niveaux hormonaux"],
        ["Infertility", "binaire", "0 / 1", "Infertilité signalée"],
        ["BMI", "continu", "15–37", "Indice de masse corporelle"],
        ["Diagnosis", "binaire", "0 / 1", "Cible — endométriose"],
    ],
)

add_heading("2.2 Statistiques descriptives", level=2)
add_bullet("10 000 observations, 7 variables")
add_bullet("Aucune valeur manquante, aucun doublon")
add_bullet("Âge moyen : 33,7 ans (écart-type 9,2)")
add_bullet("BMI moyen : 23,1 (écart-type 3,9)")
add_bullet(
    "70 % des patientes présentent une irrégularité menstruelle, 57 % une anomalie "
    "hormonale, 30 % une infertilité"
)

# ================================================================
# 3. EDA
# ================================================================
add_heading("3. Analyse exploratoire", level=1)

add_heading("3.1 Distribution de la variable cible", level=2)
add_para(
    "La variable cible est modérément déséquilibrée : 59,2 % de patientes non "
    "diagnostiquées contre 40,8 % diagnostiquées. Ce déséquilibre est pris en compte via "
    "un split stratifié et l'option class_weight='balanced'."
)
add_figure("01_target_distribution.png", width_cm=15, caption="Figure 1 — Répartition du diagnostic")

add_heading("3.2 Variables numériques", level=2)
add_para(
    "Les distributions des variables Age, Chronic_Pain_Level et BMI sont approximativement "
    "uniformes ou normales sur leurs plages respectives."
)
add_figure("02_numerical_distributions.png", width_cm=16, caption="Figure 2 — Distributions des variables numériques")
add_para(
    "Une analyse bivariée par diagnostic révèle un décalage faible mais présent pour la "
    "douleur chronique : la médiane passe de 4,85 chez les cas négatifs à 5,24 chez les "
    "cas positifs."
)
add_figure("03_boxplots_by_diagnosis.png", width_cm=16, caption="Figure 3 — Variables numériques par diagnostic")

add_heading("3.3 Variables binaires", level=2)
add_para("Toutes les variables binaires présentent une prévalence plus élevée chez les cas positifs :")
add_table(
    ["Variable", "% positif si Diagnosis=0", "% positif si Diagnosis=1", "Écart"],
    [
        ["Menstrual_Irregularity", "66,1 %", "75,0 %", "+8,9 pts"],
        ["Hormone_Level_Abnormality", "51,5 %", "70,2 %", "+18,7 pts"],
        ["Infertility", "26,2 %", "35,1 %", "+8,9 pts"],
    ],
)
add_figure("04_binary_vs_diagnosis.png", width_cm=16, caption="Figure 4 — Variables binaires vs diagnostic")

add_heading("3.4 Matrice de corrélation", level=2)
add_figure("05_correlation_matrix.png", width_cm=13, caption="Figure 5 — Matrice de corrélation")
add_para(
    "Les corrélations entre variables explicatives sont très faibles (< 0,02), ce qui "
    "suggère l'absence de multicolinéarité problématique."
)
add_figure("06_target_correlations.png", width_cm=14, caption="Figure 6 — Corrélation des variables avec la cible")
add_para("Classement des variables par corrélation avec Diagnosis :", bold=True)
add_table(
    ["Variable", "Corrélation (Pearson)"],
    [
        ["Hormone_Level_Abnormality", "0,187"],
        ["Chronic_Pain_Level", "0,117"],
        ["Infertility", "0,096"],
        ["Menstrual_Irregularity", "0,095"],
        ["BMI", "0,080"],
        ["Age", "−0,012"],
    ],
)

# ================================================================
# 4. MODÉLISATION
# ================================================================
add_heading("4. Modélisation", level=1)

add_heading("4.1 Protocole", level=2)
add_bullet("Split : 80 % entraînement / 20 % test, stratifié sur Diagnosis (random_state=42)")
add_bullet("Prétraitement : StandardScaler dans tous les pipelines")
add_bullet("Validation : Stratified K-Fold 5-plis sur le F1-score")
add_bullet("Évaluation finale : Accuracy, Precision, Recall, F1, ROC-AUC")
add_bullet("Déséquilibre : class_weight='balanced' pour LogReg, SVM et Random Forest")

add_heading("4.2 Modèles comparés", level=2)
add_table(
    ["#", "Modèle", "Hyperparamètres clés"],
    [
        ["1", "Logistic Regression", "max_iter=1000, class_weight='balanced'"],
        ["2", "K-Nearest Neighbors", "n_neighbors=7"],
        ["3", "SVM", "kernel=RBF, class_weight='balanced', probability=True"],
        ["4", "Random Forest", "n_estimators=200, class_weight='balanced'"],
        ["5", "Gradient Boosting", "n_estimators=200"],
    ],
)

add_heading("4.3 Validation croisée (5-fold)", level=2)
add_figure("07_cv_f1_scores.png", width_cm=14, caption="Figure 7 — F1-score en validation croisée")
add_para(
    "La dispersion des scores reste faible d'un pli à l'autre, ce qui indique des "
    "modèles stables."
)

add_heading("4.4 Résultats sur le test set", level=2)
add_table(
    ["Modèle", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"],
    [
        ["Logistic Regression", "0,6095", "0,5172", "0,6434", "0,5735", "0,6575"],
        ["KNN", "0,5945", "0,5035", "0,4363", "0,4675", "0,6003"],
        ["SVM", "0,6075", "0,5149", "0,6569", "0,5773", "0,6422"],
        ["Random Forest", "0,6030", "0,5159", "0,4375", "0,4735", "0,6032"],
        ["Gradient Boosting", "0,6195", "0,5518", "0,3591", "0,4350", "0,6380"],
    ],
)
add_figure("08_metric_comparison.png", width_cm=16, caption="Figure 8 — Comparaison des métriques par modèle")
add_para("Lecture :", bold=True)
add_bullet("SVM obtient le meilleur F1 grâce à un excellent rappel (65,7 %).")
add_bullet(
    "Logistic Regression est très proche sur F1 (0,574) et meilleur sur ROC-AUC (0,658) — "
    "un bon compromis performance / interprétabilité."
)
add_bullet(
    "Gradient Boosting a la meilleure accuracy, mais un rappel faible (36 %) : il manque "
    "beaucoup de cas positifs."
)
add_bullet("KNN et Random Forest sont en retrait, avec des rappels faibles.")

add_heading("4.5 Courbes ROC", level=2)
add_figure("09_roc_curves.png", width_cm=13, caption="Figure 9 — Courbes ROC")
add_para(
    "La hiérarchie des AUC confirme la robustesse de la Logistic Regression comme modèle "
    "de référence."
)

add_heading("4.6 Matrices de confusion", level=2)
add_figure("10_confusion_matrices.png", width_cm=16, caption="Figure 10 — Matrices de confusion")
add_para(
    "Les modèles équilibrés (LogReg, SVM) répartissent mieux les erreurs. Les modèles non "
    "équilibrés (GB, KNN, RF) sont biaisés vers la classe majoritaire, réduisant la "
    "détection des cas positifs."
)

add_heading("4.7 Importance des variables", level=2)
add_figure("11_feature_importance.png", width_cm=16, caption="Figure 11 — Importance des variables (modèles tree-based)")
add_para(
    "Les deux modèles tree-based confirment l'importance de Hormone_Level_Abnormality et "
    "Chronic_Pain_Level. L'âge et le BMI contribuent également (bien que faiblement "
    "corrélés linéairement), ce qui suggère une information capturée via des interactions "
    "non linéaires."
)

# ================================================================
# 5. DISCUSSION
# ================================================================
add_heading("5. Discussion", level=1)

add_heading("5.1 Pourquoi des performances modérées ?", level=2)
add_para(
    "Les scores plafonnent autour de 0,58 de F1 et 0,66 de ROC-AUC. Plusieurs raisons :"
)
add_numbered(
    "Nature synthétique du dataset : la relation entre variables et cible inclut une part "
    "significative de bruit aléatoire."
)
add_numbered("Faible nombre de variables (6 features) : peu de signal disponible.")
add_numbered(
    "Corrélations linéaires faibles : même la variable la plus corrélée (hormones) reste "
    "à r = 0,19."
)
add_numbered(
    "Absence de variables cliniques importantes : antécédents familiaux, type de douleur, "
    "imagerie, etc."
)

add_heading("5.2 Choix du modèle selon le contexte", level=2)
add_table(
    ["Objectif", "Modèle recommandé"],
    [
        ["Maximiser le rappel (dépistage)", "SVM ou Logistic Regression"],
        ["Interprétabilité (coefficients lisibles)", "Logistic Regression"],
        ["Meilleur compromis global", "Logistic Regression"],
    ],
)
add_para(
    "Dans un contexte médical de dépistage, un rappel élevé est crucial : un faux négatif "
    "(patiente malade classée saine) est plus coûteux qu'un faux positif. La Logistic "
    "Regression combine cet avantage à une interprétabilité qui facilite l'acceptation "
    "clinique."
)

# ================================================================
# 6. PISTES D'AMÉLIORATION
# ================================================================
add_heading("6. Pistes d'amélioration", level=1)
add_numbered("Tuning d'hyperparamètres : GridSearchCV ou RandomizedSearchCV.")
add_numbered("Ajustement du seuil de décision pour privilégier le rappel.")
add_numbered("Feature engineering : interactions (Age × Pain, BMI × Hormones), binning, polynômes.")
add_numbered("Modèles avancés : XGBoost, LightGBM, CatBoost, stacking.")
add_numbered("Techniques de rééquilibrage : SMOTE, undersampling, scale_pos_weight.")
add_numbered("Interprétabilité approfondie : SHAP values.")
add_numbered("Validation externe sur un dataset réel si disponible.")

# ================================================================
# 7. REPRODUCTIBILITÉ
# ================================================================
add_heading("7. Reproductibilité", level=1)
add_para("Installation et exécution :", bold=True)
code_block = (
    "git clone https://github.com/kypo-studio/endometriosis-classification.git\n"
    "cd endometriosis-classification\n"
    "uv sync\n"
    "uv run kaggle datasets download -d michaelanietie/endometriosis-dataset -p data --unzip\n"
    "uv run python reports/generate_figures.py\n"
    "uv run jupyter lab"
)
p = doc.add_paragraph()
r = p.add_run(code_block)
r.font.name = "Menlo"
r.font.size = Pt(9)

add_para("Fichiers de référence :", bold=True)
add_bullet("notebooks/01_eda.ipynb — exploration détaillée")
add_bullet("notebooks/02_modeling.ipynb — pipeline de modélisation")
add_bullet("reports/generate_figures.py — script de régénération des figures")
add_bullet("reports/model_results.csv — tableau brut des résultats")

# ================================================================
# 8. CONCLUSION
# ================================================================
add_heading("8. Conclusion", level=1)
add_para(
    "Ce projet montre qu'avec 6 variables cliniques simples, il est possible d'atteindre "
    "un pouvoir discriminant modéré (AUC ≈ 0,66) pour prédire l'endométriose sur un "
    "dataset synthétique. La Logistic Regression se dégage comme le meilleur compromis "
    "entre performance, robustesse et interprétabilité. Les gains les plus prometteurs "
    "viendraient de l'ajout de variables cliniques complémentaires et d'un tuning fin "
    "des hyperparamètres."
)

doc.save(OUT)
print(f"✓ Rapport généré : {OUT} ({OUT.stat().st_size // 1024} KB)")
